"""Document processing tools for the Processing Agent.

Python port of doc-find-replace JS logic using python-docx.
Handles find/replace, redline generation, and term extraction.
"""

import copy
import re
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from lxml import etree

NSMAP = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def read_docx_text(docx_path):
    """Extract all paragraph text from a .docx file.

    Args:
        docx_path: Path to the .docx file.

    Returns:
        Full text with paragraphs separated by newlines.
    """
    doc = Document(docx_path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _detect_case(text):
    """Detect the case pattern of a string.

    Args:
        text: The string to analyze.

    Returns:
        One of 'upper', 'lower', 'title', 'capitalized', 'mixed'.
    """
    if text == text.upper() and text != text.lower():
        return "upper"
    if text == text.lower():
        return "lower"
    words = text.split()
    if len(words) > 1 and all(w[0].isupper() for w in words if w):
        return "title"
    if text[0].isupper() and text[1:] == text[1:].lower():
        return "capitalized"
    return "mixed"


def _apply_case(replacement, case_type):
    """Apply a case pattern to replacement text.

    Args:
        replacement: The replacement string.
        case_type: One of 'upper', 'lower', 'title', 'capitalized', 'mixed'.

    Returns:
        Case-adjusted replacement string.
    """
    if case_type == "upper":
        return replacement.upper()
    if case_type == "lower":
        return replacement.lower()
    if case_type == "title":
        return replacement.title()
    if case_type == "capitalized":
        return replacement[0].upper() + replacement[1:]
    return replacement


def apply_replacements(docx_path, replacements, output_path):
    """Apply find/replace pairs to a .docx file (clean version).

    Args:
        docx_path: Path to the source .docx file.
        replacements: List of dicts with 'find' and 'replace' keys.
        output_path: Path to write the modified .docx.

    Returns:
        Dict with 'output_path' and 'change_count'.
    """
    doc = Document(docx_path)
    total_changes = 0

    for para in doc.paragraphs:
        full_text = para.text
        modified = full_text

        for r in replacements:
            pattern = re.compile(re.escape(r["find"]), re.IGNORECASE)
            def _replacer(m):
                case_type = _detect_case(m.group(0))
                return _apply_case(r["replace"], case_type)
            modified = pattern.sub(_replacer, modified)

        if modified != full_text:
            changes = sum(
                1 for _ in re.finditer(
                    re.escape(replacements[0]["find"]), full_text, re.IGNORECASE
                )
            ) if replacements else 0
            total_changes += changes

            runs = para.runs
            if runs:
                runs[0].text = modified
                for run in runs[1:]:
                    run.text = ""

    doc.save(output_path)
    return {"output_path": str(output_path), "change_count": total_changes}


def generate_redline(docx_path, replacements, output_path):
    """Create a tracked-changes version of a .docx file.

    Uses w:del and w:ins XML elements to mark deletions and insertions,
    matching the JS redline logic from doc-find-replace.

    Args:
        docx_path: Path to the source .docx file.
        replacements: List of dicts with 'find' and 'replace' keys.
        output_path: Path to write the redlined .docx.

    Returns:
        Dict with 'output_path' and 'change_count'.
    """
    doc = Document(docx_path)
    change_id = 100
    date_str = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
    author = "Clause Remediation App"
    total_changes = 0

    for para in doc.paragraphs:
        full_text = para.text
        if not full_text.strip():
            continue

        all_matches = []
        for r in replacements:
            for m in re.finditer(re.escape(r["find"]), full_text, re.IGNORECASE):
                case_type = _detect_case(m.group(0))
                all_matches.append({
                    "start": m.start(),
                    "end": m.end(),
                    "original": m.group(0),
                    "replacement": _apply_case(r["replace"], case_type),
                })

        if not all_matches:
            continue

        all_matches.sort(key=lambda x: x["start"])
        total_changes += len(all_matches)

        rpr_source = None
        for run in para.runs:
            rpr_el = run._element.find(f"{W}rPr")
            if rpr_el is not None:
                rpr_source = copy.deepcopy(rpr_el)
                break

        p_element = para._element
        for run in list(para.runs):
            p_element.remove(run._element)

        cursor = 0
        for match in all_matches:
            if cursor < match["start"]:
                _add_run(p_element, full_text[cursor:match["start"]], rpr_source)

            del_el = etree.SubElement(p_element, f"{W}del")
            del_el.set(f"{W}id", str(change_id))
            del_el.set(f"{W}author", author)
            del_el.set(f"{W}date", date_str)
            change_id += 1

            del_run = etree.SubElement(del_el, f"{W}r")
            if rpr_source is not None:
                del_run.append(copy.deepcopy(rpr_source))
            del_text = etree.SubElement(del_run, f"{W}delText")
            del_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            del_text.text = match["original"]

            ins_el = etree.SubElement(p_element, f"{W}ins")
            ins_el.set(f"{W}id", str(change_id))
            ins_el.set(f"{W}author", author)
            ins_el.set(f"{W}date", date_str)
            change_id += 1

            ins_run = etree.SubElement(ins_el, f"{W}r")
            if rpr_source is not None:
                ins_run.append(copy.deepcopy(rpr_source))
            ins_text = etree.SubElement(ins_run, f"{W}t")
            ins_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            ins_text.text = match["replacement"]

            cursor = match["end"]

        if cursor < len(full_text):
            _add_run(p_element, full_text[cursor:], rpr_source)

    doc.save(output_path)
    return {"output_path": str(output_path), "change_count": total_changes}


def _add_run(p_element, text, rpr_source):
    """Add a plain text run to a paragraph element.

    Args:
        p_element: The w:p lxml element.
        text: Text content for the run.
        rpr_source: Optional run properties element to clone.
    """
    run = etree.SubElement(p_element, f"{W}r")
    if rpr_source is not None:
        run.append(copy.deepcopy(rpr_source))
    t = etree.SubElement(run, f"{W}t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text


def extract_defined_terms(docx_path):
    """Extract defined terms from a legal document.

    Finds bracketed [terms], quoted "Terms", and capitalized Defined Terms.

    Args:
        docx_path: Path to the .docx file.

    Returns:
        List of dicts with 'term', 'type', and 'count' keys.
    """
    text = read_docx_text(docx_path)
    results = []
    seen = set()

    for m in re.finditer(r'\[([a-zA-Z0-9][a-zA-Z0-9 -]*[a-zA-Z0-9])\]', text):
        key = m.group(1).lower()
        if key not in seen:
            seen.add(key)
            count = len(re.findall(re.escape(m.group(0)), text, re.IGNORECASE))
            results.append({"term": m.group(1), "type": "bracket", "count": count})

    for m in re.finditer(r'["“”]([A-Z][^"“”\n]{0,79})["“”]', text):
        term = m.group(1).strip()
        key = term.lower()
        if key not in seen and len(term) >= 2:
            seen.add(key)
            count = len(re.findall(r'\b' + re.escape(term) + r'\b', text, re.IGNORECASE))
            results.append({"term": term, "type": "quoted", "count": count})

    cap_pattern = re.compile(
        r'\b([A-Z][a-z]+(?:-[A-Z][a-z]+)*'
        r'(?:\s+(?:of|the|and|or|for|in|on|to|by|a|an|at|as|with|from)'
        r'\s+[A-Z][a-z]+(?:-[A-Z][a-z]+)*'
        r'|\s+[A-Z][a-z]+(?:-[A-Z][a-z]+)*)+)\b'
    )
    cap_counts = {}
    for m in cap_pattern.finditer(text):
        phrase = re.sub(r'^(?:The|A|An)\s+', '', m.group(1))
        key = phrase.lower()
        if key in seen:
            continue
        cap_counts[key] = cap_counts.get(key, {"phrase": phrase, "count": 0})
        cap_counts[key]["count"] += 1

    for key, info in cap_counts.items():
        if info["count"] >= 2:
            seen.add(key)
            results.append({"term": info["phrase"], "type": "defined", "count": info["count"]})

    return results
